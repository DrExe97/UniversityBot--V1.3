"""
document_loader.py - Extract text from PDF, DOCX, and TXT files
Prepares document content for the knowledge base
"""

import os
import re
import logging
from pathlib import Path

# ── Import ALL libraries at module level (not inside functions)
# Lazy imports inside functions fail in background threads on Windows

try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc", ".txt", ".md"}


def is_supported(filename: str) -> bool:
    ext = Path(filename).suffix.lower()
    return ext in SUPPORTED_EXTENSIONS


def extract_from_pdf(filepath: str) -> str:
    if not PYPDF2_AVAILABLE:
        raise ImportError("PyPDF2 not installed. Run: pip install PyPDF2")
    text_parts = []
    with open(filepath, "rb") as f:
        reader = PyPDF2.PdfReader(f)
        total_pages = len(reader.pages)
        for page_num, page in enumerate(reader.pages):
            try:
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    text_parts.append(f"[Page {page_num + 1}/{total_pages}]\n{page_text.strip()}")
            except Exception as e:
                logger.warning(f"Could not extract page {page_num + 1}: {e}")
    full_text = "\n\n".join(text_parts)
    logger.info(f"Extracted {len(full_text)} chars from PDF ({total_pages} pages)")
    return full_text


def extract_from_docx(filepath: str) -> str:
    try:
        import docx as python_docx
    except ImportError:
        raise ImportError("python-docx not installed. Run: pip install python-docx")
    doc = python_docx.Document(filepath)
    text_parts = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            if para.style.name.startswith("Heading"):
                text_parts.append(f"\n## {text}\n")
            else:
                text_parts.append(text)
    for table in doc.tables:
        table_text = []
        for row in table.rows:
            row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_cells:
                table_text.append(" | ".join(row_cells))
        if table_text:
            text_parts.append("\n[Table]\n" + "\n".join(table_text) + "\n")
    full_text = "\n".join(text_parts)
    logger.info(f"Extracted {len(full_text)} chars from DOCX")
    return full_text


def extract_from_txt(filepath: str) -> str:
    encodings = ["utf-8", "utf-8-sig", "latin-1", "cp1252"]
    for encoding in encodings:
        try:
            with open(filepath, "r", encoding=encoding) as f:
                text = f.read()
            logger.info(f"Extracted {len(text)} chars from TXT (encoding: {encoding})")
            return text.strip()
        except UnicodeDecodeError:
            continue
    raise ValueError(f"Could not decode {filepath} with any supported encoding")


def extract_text(filepath: str) -> str:
    path = Path(filepath)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {filepath}")
    ext = path.suffix.lower()
    filename = path.name
    logger.info(f"Extracting text from: {filename} ({ext})")
    if ext == ".pdf":
        text = extract_from_pdf(filepath)
    elif ext in (".docx", ".doc"):
        text = extract_from_docx(filepath)
    elif ext in (".txt", ".md"):
        text = extract_from_txt(filepath)
    else:
        raise ValueError(f"Unsupported file format: {ext}. Supported: {', '.join(SUPPORTED_EXTENSIONS)}")
    text = clean_text(text)
    if not text:
        raise ValueError(f"No text could be extracted from {filename}")
    return text


def clean_text(text: str) -> str:
    if not text:
        return ""
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ \t]{2,}", " ", text)
    text = re.sub(r"[^\x20-\x7E\n\t\u00A0-\uFFFF]", "", text)
    text = re.sub(r"-\n([a-z])", r"\1", text)
    return text.strip()


def get_file_metadata(filepath: str) -> dict:
    path = Path(filepath)
    size_bytes = path.stat().st_size
    return {
        "filename":   path.name,
        "extension":  path.suffix.lower(),
        "size_bytes": size_bytes,
        "size_kb":    round(size_bytes / 1024, 2),
    }